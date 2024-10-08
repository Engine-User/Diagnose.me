import streamlit as st
import os
from langchain.tools import Tool
from langchain.utilities import GoogleSerperAPIWrapper
from dotenv import load_dotenv
from docx import Document
from io import BytesIO
import base64
from datetime import datetime, timedelta
from langchain_groq import ChatGroq
import PyPDF2
import docx2txt
from PIL import Image

# Load environment variables
load_dotenv()

# Set the API keys from environment variables
groq_api_key = os.getenv("GROQ_API_KEY")
serper_api_key = os.getenv("SERPER_API_KEY")

# Check if the API keys are set
if not groq_api_key or not serper_api_key:
    st.error("Please set the GROQ_API_KEY and SERPER_API_KEY in your .env file.")
    st.stop()

# Initialize ChatGroq
try:
    llm = ChatGroq(
        api_key=groq_api_key,
        model_name='llama-3.1-70b-versatile',  # Updated to a valid Groq model
        temperature=0.7
    )
except Exception as e:
    st.error(f"Error initializing ChatGroq: {str(e)}")
    st.stop()

# Initialize search tool
search_tool = None
try:
    search = GoogleSerperAPIWrapper(serper_api_key=serper_api_key)
    search_tool = Tool(
        name="Search",
        func=search.run,
        description="useful for when you need to answer questions about current events; need to gather information about the user inputs. Attend and search main key words and phrases"
    )
except Exception as e:
    st.error(f"Error initializing search tool: {str(e)}")
    st.stop()

# Initialize Medication Management
medications = {}

# Initialize Mental Health Support
mental_health_scores = {}

# Medication Management functions
def add_medication(name, dosage, frequency):
    medications[name] = {"dosage": dosage, "frequency": frequency, "next_dose": datetime.now()}

def get_medication_reminders():
    reminders = []
    now = datetime.now()
    for med, info in medications.items():
        if now >= info["next_dose"]:
            reminders.append(f"Time to take {med} - Dosage: {info['dosage']}")
            if info["frequency"] == "daily":
                info["next_dose"] = now + timedelta(days=1)
            elif info["frequency"] == "twice_daily":
                info["next_dose"] = now + timedelta(hours=12)
    return reminders

# Mental Health Support functions
def mental_health_screening(answers):
    score = sum(answers)
    if score < 5:
        return "Low risk", score
    elif 5 <= score < 10:
        return "Moderate risk", score
    else:
        return "High risk", score

def update_mood_tracker(mood):
    today = datetime.now().strftime("%Y-%m-%d")
    mood_message = {
        "Very Bad": "You're having a really tough day. Remember, it's okay to not be okay.",
        "Bad": "Things are a bit rough today. Take some time for self-care.",
        "Neutral": "You're feeling balanced today. That's okay too!",
        "Good": "You're having a good day! Enjoy the positive feelings.",
        "Very Good": "You're feeling great today! Celebrate the good moments."
    }
    mental_health_scores[today] = mood_message[mood]
    # Ensure we only keep the last 7 days of data
    if len(mental_health_scores) > 7:
        oldest_date = min(mental_health_scores.keys())
        del mental_health_scores[oldest_date]

def get_mood_trend():
    # Sort the mood data by date
    sorted_moods = sorted(mental_health_scores.items(), key=lambda x: x[0], reverse=True)
    # Return the moods with their dates
    return [f"{date}: {mood}" for date, mood in sorted_moods]

# Function to generate a Word document from the result
def generate_docx(result):
    doc = Document()
    doc.add_heading('Healthcare Diagnosis and Treatment Recommendations', 0)
    
    # Split the result into lines
    lines = result.split('\n')
    
    for line in lines:
        if line.startswith('# '):
            doc.add_heading(line[2:], level=1)
        elif line.startswith('## '):
            doc.add_heading(line[3:], level=2)
        elif line.startswith('*'):
            doc.add_paragraph(line, style='Intense Quote')
        else:
            doc.add_paragraph(line)
    
    bio = BytesIO()
    doc.save(bio)
    bio.seek(0)
    return bio

# Function to generate a download link for the Word document
def get_download_link(bio, filename):
    b64 = base64.b64encode(bio.read()).decode()
    return f'<a href="data:application/vnd.openxmlformats-officedocument.wordprocessingml.document;base64,{b64}" download="{filename}">Download Diagnosis and Treatment Plan</a>'

# Function to read PDF files
def read_pdf(file):
    pdf_reader = PyPDF2.PdfReader(file)
    text = ""
    for page in pdf_reader.pages:
        text += page.extract_text() + "\n"
    return text

# Function to read DOCX files
def read_docx(file):
    return docx2txt.process(file)

# Function to handle image files
def handle_image(file):
    return "An image file was uploaded. Please describe its contents in the symptoms or medical history fields if relevant."

# Streamlit UI setup
st.set_page_config(layout="wide")

# Custom CSS for styling
st.markdown("""
<style>
    @import url('https://fonts.googleapis.com/css2?family=Graduate&display=swap');

    body {
        font-family: 'Graduate', sans-serif !important;
    }

    .stButton > button {
        background-color: #1E90FF !important;  /* Changed to blue */
        color: white !important;
        font-family: 'Graduate', sans-serif !important;
        border-radius: 10px;
        border: none;
        box-shadow: 0px 0px 15px #1E90FF;  /* Changed to blue */
        transition: 0.3s;
        width: 100%;
        margin: 5px 0;
    }

    .stButton > button:hover {
        box-shadow: 0px 0px 25px 10px #1E90FF;  /* Changed to blue */
        transform: scale(1.05);
    }

    .title-block {
        background-color: #1E90FF;
        color: white;
        padding: 20px;
        border-radius: 8px;
        box-shadow: 4px 4px 40px #1E90FF;
        text-align: center;
        transition: all 0.3s ease-in-out;
        margin-bottom: 20px;
    }

    .title-block:hover {
        box-shadow: 0px 0px 30px 10px #1E90FF;
        transform: scale(1.05);
    }

    .title-block h1 {
        color: white !important;
        margin: 0;
        padding: 0;
        font-size: 2.5em;
        text-shadow: 2px 2px 4px rgba(0,0,0,0.5);
    }

    h1, h2, h3, h4, h5, h6 {
        color: #1E90FF !important;  /* Changed to blue */
        font-family: 'Graduate', sans-serif !important;
    }

    .stSelectbox, .stNumberInput, .stTextArea, .stFileUploader {
        font-family: 'Graduate', sans-serif !important;
    }

    .stSelectbox > div > div > div {
        font-family: 'Graduate', sans-serif !important;
    }

    .stTextArea textarea {
        font-family: 'Graduate', sans-serif !important;
    }

    .stFileUploader > div > div > div {
        font-family: 'Graduate', sans-serif !important;
    }

    /* Sidebar styling */
    .sidebar-content {
        padding: 20px;
    }

    .sidebar-button {
        margin-bottom: 15px;
    }
</style>
""", unsafe_allow_html=True)

# Main title in a blue block
st.markdown('<div class="title-block"><h1>Medical Diagnosis and Treatment Plan</h1></div>', unsafe_allow_html=True)

# Sidebar
with st.sidebar:
    st.markdown("""
    <div style="font-family: 'Graduate', sans-serif; background-color: #000080; color: white; padding: 20px; border-radius: 10px; border: 2px solid #1E90FF;">
        <h2 style="color: white; font-size: 3.0em;">Diagnose.Me</h2>
    </div>
    """, unsafe_allow_html=True)
    st.markdown('<div class="sidebar-content">', unsafe_allow_html=True)
    
    st.markdown('<div class="sidebar-button">', unsafe_allow_html=True)
    about_button = st.button("About")
    st.markdown('</div>', unsafe_allow_html=True)
    
    st.markdown('<div class="sidebar-button">', unsafe_allow_html=True)
    contact_button = st.button("Contact")
    st.markdown('</div>', unsafe_allow_html=True)
    
    if about_button:
        st.markdown("""
        <div style="font-family: 'Graduate', sans-serif; background-color: #000080; color: white; padding: 20px; border-radius: 10px; border: 2px solid #1E90FF;">
            <h3 style="color: #ADD8E6;">About DiagnoseMe</h3>
            <p>**Welcome to DiagnoseMe**</p>
            <ul>
                <li>In today's digital age, accessing accurate and reliable health information has become increasingly challenging. With Diagnose.Me, you can now enjoy peace of mind, focusing on what matters most while our AI-powered platform serves as your trusted, personal doctor, assessing your symptoms and cross-checking with vast amounts of medical data to provide personalized guidance and support.</li>    
                <li>Our algorithm uses multiple AI agents working in tandem to aggregate real-time information related to your condition, ensuring you receive the most up-to-date and effective diagnosis and recommended course of action. By attaching relevant files and test results, our AI agents can better understand your unique medical history, making informed decisions tailored to your needs..</li>
                <li>Our platform uses advanced AI technology to analyze your symptoms, medical history, and other relevant health data to provide you with personalized diagnostic insights. Our algorithm is trained on a vast repository of medical data and is constantly updated to ensure that our insights are always accurate and up-to-date.</li>
                <li>At Diagnose.Me, we believe health and wellness are best managed through peer-to-peer support and accessible, reliable information. We're committed to providing you with a trusted tool for personal diagnosis, empowering you to take control of your health and make informed decisions about your well-being..</li>
            </ul>
            <p><strong>Important:</strong> Always consult with a healthcare professional for accurate medical advice.</p>
        </div>
        """, unsafe_allow_html=True)
    
    if contact_button:
        st.markdown("""
        <div style="font-family: 'Graduate', sans-serif; background-color: #000080; color: white; padding: 20px; border-radius: 10px; border: 2px solid #1E90FF;">
            <h3 style="color: #ADD8E6;">Contact Information</h3>
            <p>Made by Engineer <a href="mailto:ggengineerco@gmail.com">ggengineerco@gmail.com</a></p>
        </div>
        """, unsafe_allow_html=True)
    

# User inputs
st.markdown('<h2>Patient Information</h2>', unsafe_allow_html=True)
gender = st.selectbox('Select Gender', ('Male', 'Female', 'Other'))
age = st.number_input('Enter Age', min_value=0, max_value=120, value=25)
symptoms = st.text_area('Enter Symptoms', '(eg:- Fever, Cold etc)')
medical_history = st.text_area('Enter Medical History', 'e.g:- diabetes, hypertension, or any other conditions from your past')

# File upload
st.markdown('<h2>Additional Documents</h2>', unsafe_allow_html=True)
uploaded_file = st.file_uploader(
    "Upload any other documents explaining your condition(Prescriptions, Lab results, Old MRIs, CT Scans, etc.)", 
    type=['pdf', 'docx', 'png', 'jpg', 'jpeg']
)
file_content = ""
if uploaded_file is not None:
    if uploaded_file.type == "application/pdf":
        file_content = read_pdf(uploaded_file)
    elif uploaded_file.type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
        file_content = read_docx(uploaded_file)
    elif uploaded_file.type in ["image/png", "image/jpeg", "image/jpg"]:
        file_content = handle_image(uploaded_file)
    st.success("File uploaded successfully!")

# Medication Management
st.markdown('<h2>Medication Management</h2>', unsafe_allow_html=True)
med_name = st.text_input("Medication Name")
med_dosage = st.text_input("Dosage")
med_frequency = st.selectbox("Frequency", ["Daily", "Twice_Daily"])
if st.button("Add Medication"):
    add_medication(med_name, med_dosage, med_frequency)
    st.success(f"Added {med_name} to your medications.")

reminders = get_medication_reminders()
if reminders:
    st.warning("Medication Reminders:")
    for reminder in reminders:
        st.write(reminder)

# Mental Health Support
st.markdown('<h2>Mental Health Support</h2>', unsafe_allow_html=True)
st.write("Please answer the following questions on a scale of 0-5 (0=Not at all & 5=All the time)")
q1 = st.slider("Interest or pleasure in doing things", 0, 5, 0)
q2 = st.slider("Feelings of despair, depression, or hopelessness", 0, 5, 0)
q3 = st.slider("Trouble falling or staying asleep, or sleeping too much", 0, 5, 0)
q4 = st.slider("Feelings of tiredness or having little energy", 0, 5, 0)
q5 = st.slider("Poor appetite or overeating", 0, 5, 0)

if st.button("Submit Mental Health Screening"):
    risk_level, score = mental_health_screening([q1, q2, q3, q4, q5])
    st.write(f"Your mental health risk level: {risk_level} (Score: {score})")
    if risk_level == "High risk":
        st.warning("You should consider consulting a mental health professional.")

mood = st.selectbox("How are you feeling today?", ["Very Bad", "Bad", "Neutral", "Good", "Very Good"])
if st.button("Update Mood"):
    update_mood_tracker(mood)
    st.success(f"Mood updated successfully. {mental_health_scores[datetime.now().strftime('%Y-%m-%d')]}")

# Function to extract and format the result
def format_result(diagnosis_and_treatment):
    # Get medication reminders
    med_reminders = get_medication_reminders()
    if med_reminders:
        medication_text = "\n".join(med_reminders)
    else:
        medication_text = "No medication reminders at this time."
    
    # Get mental health trend
    mood_trend = get_mood_trend()
    if mood_trend:
        mood_text = "\n".join(mood_trend)
    else:
        mood_text = "No mood data available for the last 7 days. Start tracking your mood to see trends!"
    
    formatted_result = f"""
# Diagnosis and Treatment Plan

{diagnosis_and_treatment}

## Medication Reminders
{medication_text}

## Mental Health Trend (Last 7 days)
{mood_text}

---
*Note: This is just the start, if you are concerned about the diagnosis or treatment plan, please consult with a healthcare professional for accurate diagnosis and treatment.*
    """
    return formatted_result

# Execution button for generating diagnosis and treatment plan
if st.button("Get Diagnosis and Treatment Plan"):
    with st.spinner('Diagnosing...'):
        try:
            # Combine all input into a single prompt
            prompt = f"""
            Analyze the following patient information and provide a diagnosis and treatment plan:
            
            Gender: {gender}
            Age: {age}
            Symptoms: {symptoms}
            Medical History: {medical_history}
            Additional Information: {file_content}
            
            Please provide:
            1. A detailed, step-by-step, effective, scientifically proven preliminary diagnosis.
            2. A comprehensive treatment plan tailored to the patient's diagnosis, symptoms, and medical history.
            
            Format your response with clear headings and bullet points for easy readability.
            """
            
            # Get response from ChatGroq
            response = llm.predict(prompt)
            
            # Format the result
            formatted_result = format_result(response)
            
            st.markdown(formatted_result)
            
            docx_file = generate_docx(formatted_result)
            download_link = get_download_link(docx_file, "diagnosis_and_treatment_plan.docx")
            st.markdown(download_link, unsafe_allow_html=True)

        except Exception as e:
            st.error(f"An error occurred: {str(e)}")

# Add this at the end of your script to ensure all UI elements are rendered
if __name__ == "__main__":
    st.sidebar.markdown("---")
    st.sidebar.text("© 2024 Diagnose.Me")
